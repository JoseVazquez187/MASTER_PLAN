from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Crear documento Word
doc = Document()

# Título
title = doc.add_paragraph("Checklist de Validación MRP - KPI's Previos a la Corrida")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.runs[0]
run.font.size = Pt(16)
run.bold = True

doc.add_paragraph("\nEste documento debe ser revisado y firmado por los responsables antes de considerar una corrida de MRP como la más limpia.\n")

# Lista de KPI's
kpis = [
    ("01 FCST VS PP", "Todo lo que esté cargado en FCST debe estar en un plan de producción."),
    ("02 PLAN DE LIBERACIONES", "Las liberaciones de las WO deben estar conforme a necesidad de intermedios."),
    ("03 ALINEACIÓN DE FECHAS EN WO VS (FCST, SO)", "Toda WO liberada que su demanda sea FCST o SO debe estar alineada con la fecha de necesidad."),
    ("04 LIMPIEZA DE PSU", "La línea PSU requiere una limpieza manual, cada acción consta de 4 facturables, solo dos se cierran automáticamente."),
    ("05 LIMPIEZA DE P&S DE AC CON TERMINACIÓN 'B'", "El material P&S de los aviones con terminación B es necesario cerrarlos manualmente."),
    ("06 CREDIT MEMOS ABIERTOS", "Un credit memo abierto sin surtir es un potencial para duplicar demanda."),
    ("07 WO EN FIRME Y SIN EXPLOSIÓN DE MATERIALES", "Validación de materiales en WO Released."),
    ("08 REDUCCIÓN DE ÓRDENES A CANCELAR", "Eliminar órdenes a cancelar."),
    ("09 WO ACORDE A REVISIÓN EN R4", "Toda WO liberada debe de estar en la revisión más actual de R4."),
    ("10 NÚMEROS OBSOLETOS/EXPIRADOS EN LOCALIDAD NETEABLE", "Mover números obsoletos a localidades no neteables.")
]

for kpi, desc in kpis:
    p = doc.add_paragraph()
    p.add_run(f"{kpi}: ").bold = True
    p.add_run(desc)

# Espacio para firmas
doc.add_paragraph("\n\nValidado por:\n")
doc.add_paragraph("_________________________          _________________________")
doc.add_paragraph("Nombre y Firma Responsable          Nombre y Firma Manager")

# Guardar archivo
output_path = "/mnt/data/Checklist_Validacion_MRP.docx"
doc.save(output_path)

output_path
